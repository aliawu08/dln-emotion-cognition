#!/usr/bin/env python3
"""Convert markdown rubric/questionnaire files to .docx format."""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

PROTOCOL_DIR = Path(__file__).parent


def parse_md_and_build_docx(md_path: Path, docx_path: Path):
    """Parse a markdown file and produce a formatted .docx."""
    text = md_path.read_text()
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Horizontal rule — skip
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Clean markdown bold/italic from headings
            heading_text = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", heading_text)
            p = doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[-\s|:]+\|$", lines[i + 1].strip()):
            # Parse table header
            headers = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator

            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1

            # Create table
            table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Light Grid Accent 1")
            # Header row
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
            for ri, row in enumerate(rows):
                for j, cell in enumerate(row):
                    if j < len(table.columns):
                        table.rows[ri + 1].cells[j].text = cell
            doc.add_paragraph()  # spacing after table
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Numbered list item
        num_match = re.match(r"^(\d+)\.\s+(.*)", line)
        if num_match:
            item_text = num_match.group(2)
            add_rich_paragraph(doc, item_text, style="List Number")
            i += 1
            continue

        # Bullet list item (including [ ] checkboxes)
        bullet_match = re.match(r"^[\-\*]\s+(.*)", line.strip())
        if bullet_match:
            item_text = bullet_match.group(1)
            add_rich_paragraph(doc, item_text, style="List Bullet")
            i += 1
            continue

        # Regular paragraph (may span multiple lines)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != "" and not lines[i].startswith("#") and not lines[i].startswith("|") and lines[i].strip() not in ("---", "***", "___"):
            if re.match(r"^[\-\*]\s+", lines[i].strip()) or re.match(r"^\d+\.\s+", lines[i]):
                break
            para_lines.append(lines[i])
            i += 1

        full_text = " ".join(l.strip() for l in para_lines)
        add_rich_paragraph(doc, full_text)
        continue

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


def add_rich_paragraph(doc, text, style=None):
    """Add a paragraph with bold/italic markdown formatting handled."""
    p = doc.add_paragraph(style=style)
    # Split on bold (**...**) and italic (*...*) markers
    # Process bold first, then italic within segments
    segments = re.split(r"(\*{1,2}.*?\*{1,2})", text)
    for seg in segments:
        if seg.startswith("**") and seg.endswith("**"):
            run = p.add_run(seg[2:-2])
            run.bold = True
        elif seg.startswith("*") and seg.endswith("*"):
            run = p.add_run(seg[1:-1])
            run.italic = True
        else:
            # Handle inline code
            code_parts = re.split(r"(`[^`]+`)", seg)
            for part in code_parts:
                if part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                else:
                    # Handle underline blanks
                    p.add_run(part)


if __name__ == "__main__":
    files_to_convert = [
        "coder_screening_questionnaire.md",
        "blinded_coding_rubric_webb.md",
        "blinded_coding_rubric_hoyt.md",
    ]

    for fname in files_to_convert:
        md_path = PROTOCOL_DIR / fname
        docx_path = PROTOCOL_DIR / fname.replace(".md", ".docx")
        if md_path.exists():
            parse_md_and_build_docx(md_path, docx_path)
        else:
            print(f"Not found: {md_path}")
